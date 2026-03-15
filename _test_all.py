"""Комплексный тест всех парсеров на случайных тестовых файлах."""

import pathlib
import random
import shutil
import string
import tempfile

from docx import Document as DocxDocument

from parser_manager.models import CorruptedFileError, ParsingFailedError
from parser_manager.parsers.documents.djvu_parser import DjvuParser
from parser_manager.parsers.documents.doc_parser import DocParser
from parser_manager.parsers.documents.docx_parser import DocxParser
from parser_manager.parsers.documents.pdf_parser import PdfParser
from parser_manager.parsers.html_parser import HtmlParser

# ── helpers ──────────────────────────────────────────────────────────────────


def rand_text(n=6):
    return "".join(random.choices(string.ascii_letters + " ", k=n)).strip()


OK = "\033[32mOK\033[0m"
ERR = "\033[31mFAIL\033[0m"
results = []


def check(label, condition, got=""):
    status = OK if condition else ERR
    msg = f"  [{status}] {label}"
    if not condition:
        msg += f"  → got: {got!r}"
    print(msg)
    results.append(condition)


# ── HTML ─────────────────────────────────────────────────────────────────────

title = rand_text(10)
heading = rand_text(8)
para = rand_text(20)
author = rand_text(8)

html_content = f"""<!DOCTYPE html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="author" content="{author}">
  <title>{title}</title>
</head>
<body>
  <h1>{heading}</h1>
  <p>{para}</p>
  <a href="https://example.com">link</a>
</body>
</html>"""

f_html = pathlib.Path(tempfile.mktemp(suffix=".html"))
f_html.write_text(html_content, encoding="utf-8")

print("\n=== HTML ===")
r = HtmlParser(str(f_html)).parse()
check("success=True", r.success)
check("format='html'", r.format == "html")
check("title в metadata", r.metadata.get("title") == title, r.metadata.get("title"))
check("author в metadata", r.metadata.get("author") == author, r.metadata.get("author"))
check("heading в structure", any(e["element_type"] == "heading" for e in r.structure))
check("paragraph в structure", any(e["element_type"] == "paragraph" for e in r.structure))
check("link в structure", any(e["element_type"] == "link" for e in r.structure))
check("текст содержит heading", heading in r.text)
f_html.unlink()

# ── PDF ──────────────────────────────────────────────────────────────────────

pdf_text = rand_text(30)


def make_minimal_pdf(text: str) -> bytes:
    stream_content = f"BT /F1 12 Tf 50 700 Td ({text}) Tj ET".encode("latin-1", errors="replace")
    objects = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 5 0 R /Resources << /Font << /F1 4 0 R >> >> >>"
        ),
        4: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        5: (
            f"<< /Length {len(stream_content)} >>\nstream\n".encode()
            + stream_content
            + b"\nendstream"
        ),
    }
    buf = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for n in sorted(objects):
        offsets[n] = len(buf)
        buf += f"{n} 0 obj\n".encode() + objects[n] + b"\nendobj\n"
    xref_pos = len(buf)
    xref = f"xref\n0 {max(objects) + 1}\n0000000000 65535 f \n"
    for i in range(1, max(objects) + 1):
        xref += f"{offsets.get(i, 0):010d} 00000 {'n' if i in offsets else 'f'} \n"
    trailer = f"trailer\n<< /Size {max(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n"
    buf += xref.encode() + trailer.encode()
    return bytes(buf)


f_pdf = pathlib.Path(tempfile.mktemp(suffix=".pdf"))
f_pdf.write_bytes(make_minimal_pdf(pdf_text))

print("\n=== PDF ===")
r = PdfParser(str(f_pdf)).parse()
check("success=True", r.success)
check("format='pdf'", r.format == "pdf")
check("pages=1", r.raw_data.get("pages") == 1, r.raw_data.get("pages"))
check("текст извлечён", pdf_text in r.text, r.text[:60])
check("structure не пуст", len(r.structure) > 0)
f_pdf.unlink()

# ── DOCX ─────────────────────────────────────────────────────────────────────

docx_heading = rand_text(10)
docx_para = rand_text(25)
docx_author = rand_text(8)
docx_title = rand_text(12)

f_docx = pathlib.Path(tempfile.mktemp(suffix=".docx"))
doc = DocxDocument()
doc.add_heading(docx_heading, level=1)
doc.add_paragraph(docx_para)
t = doc.add_table(rows=1, cols=2)
t.rows[0].cells[0].text = "Cell A"
t.rows[0].cells[1].text = "Cell B"
doc.core_properties.author = docx_author
doc.core_properties.title = docx_title
doc.save(str(f_docx))

print("\n=== DOCX ===")
r = DocxParser(str(f_docx)).parse()
check("success=True", r.success)
check("format='docx'", r.format == "docx")
check("title в metadata", r.metadata.get("title") == docx_title, r.metadata.get("title"))
check(
    "author в metadata",
    r.metadata.get("author") == docx_author,
    r.metadata.get("author"),
)
check("heading в structure", any(e["element_type"] == "heading" for e in r.structure))
check("paragraph в structure", any(e["element_type"] == "paragraph" for e in r.structure))
check("table в structure", any(e["element_type"] == "table" for e in r.structure))
check("текст содержит heading", docx_heading in r.text)
f_docx.unlink()

# ── DOC: проверяем CorruptedFileError / fallback ──────────────────────────────

print("\n=== DOC (некорректный файл → CorruptedFileError) ===")
f_doc = pathlib.Path(tempfile.mktemp(suffix=".doc"))
f_doc.write_text("не OLE файл", encoding="utf-8")
try:
    DocParser(str(f_doc)).extract_metadata()
    check("CorruptedFileError поднят", False, "исключение не поднялось")
except CorruptedFileError:
    check("CorruptedFileError поднят", True)
except Exception as ex:
    check("CorruptedFileError поднят", False, type(ex).__name__)
f_doc.unlink()

# ── DJVU: проверяем ParsingFailedError при отсутствии djvutxt ────────────────

print("\n=== DJVU (нет djvutxt → ParsingFailedError) ===")
if shutil.which("djvutxt") is None:
    f_djvu = pathlib.Path(tempfile.mktemp(suffix=".djvu"))
    f_djvu.write_bytes(b"FAKE DJVU")
    try:
        DjvuParser(str(f_djvu)).extract_text()
        check("ParsingFailedError поднят", False, "исключение не поднялось")
    except ParsingFailedError:
        check("ParsingFailedError поднят", True)
    except Exception as ex:
        check("ParsingFailedError поднят", False, type(ex).__name__)
    f_djvu.unlink()
else:
    print("  [--] djvutxt найден, пропускаем проверку ошибки")

# ── Итог ─────────────────────────────────────────────────────────────────────
print(f"\n{'=' * 40}")
passed = sum(results)
total = len(results)
color = "\033[32m" if passed == total else "\033[33m"
print(f"{color}Пройдено: {passed}/{total}\033[0m")
