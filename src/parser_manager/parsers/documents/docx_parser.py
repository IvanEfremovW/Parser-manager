"""
DOCX-парсер на основе python-docx
"""

import logging

from docx import Document as load_docx_document
from docx.document import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from parser_manager.core.base_parser import BaseParser
from parser_manager.models import (
    ParsedContent,
    DocumentMetadata,
    TextElement,
    ParsingFailedError,
    CorruptedFileError,
)
from parser_manager.utils import (
    derive_semantic_blocks,
    semantic_summary,
    score_quality,
    collect_file_metrics,
)

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Парсер DOCX-файлов."""

    supported_extensions: tuple = (".docx",)
    format_name: str = "docx"

    def _load_document(self) -> DocxDocument:
        try:
            return load_docx_document(str(self.file_path))
        except PackageNotFoundError as exc:
            raise CorruptedFileError(
                f"Файл '{self.file_path.name}' не является валидным DOCX"
            ) from exc
        except Exception as exc:
            raise CorruptedFileError(
                f"Не удалось открыть DOCX '{self.file_path.name}': {exc}"
            ) from exc

    def extract_text(self) -> str:
        doc = self._load_document()
        blocks: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_cells = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_cells:
                    blocks.append(" | ".join(row_cells))

        return "\n".join(blocks).strip()

    def extract_metadata(self) -> DocumentMetadata:
        doc = self._load_document()
        props = doc.core_properties

        return DocumentMetadata(
            title=props.title or None,
            author=props.author or None,
            subject=props.subject or None,
            creation_date=props.created,
            modification_date=props.modified,
            custom_fields={
                "category": props.category,
                "comments": props.comments,
                "keywords": props.keywords,
                "last_modified_by": props.last_modified_by,
            },
        )

    def extract_structure(self) -> list:
        doc = self._load_document()
        elements: list[TextElement] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style_name.startswith("heading"):
                level = 1
                tokens = style_name.split()
                if len(tokens) > 1 and tokens[-1].isdigit():
                    level = int(tokens[-1])
                elements.append(
                    TextElement(content=text, element_type="heading", level=level)
                )
            elif "list" in style_name:
                elements.append(TextElement(content=text, element_type="list"))
            else:
                elements.append(TextElement(content=text, element_type="paragraph"))

        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                rows.append(row_text)
            table_content = "\n".join(rows).strip()
            if table_content:
                elements.append(
                    TextElement(content=table_content, element_type="table")
                )

        return [element.to_dict() for element in elements]

    def parse(self) -> ParsedContent:
        try:
            text = self.extract_text()
            metadata = self.extract_metadata()
            structure = self.extract_structure()
            semantic_blocks = derive_semantic_blocks(text, structure)
            quality = score_quality(text, semantic_blocks)
            file_metrics = collect_file_metrics(
                str(self.file_path), semantic_blocks, text
            )

            return ParsedContent(
                file_path=str(self.file_path),
                format=self.format_name,
                text=text,
                metadata=metadata.to_dict(),
                structure=structure,
                semantic_blocks=semantic_blocks,
                quality=quality,
                file_metrics=file_metrics,
                raw_data={
                    "paragraphs": len(
                        [s for s in structure if s["element_type"] == "paragraph"]
                    ),
                    "semantic_summary": semantic_summary(semantic_blocks),
                },
                success=True,
            )
        except (CorruptedFileError, ParsingFailedError):
            raise
        except Exception as exc:
            logger.exception("Ошибка при парсинге DOCX: %s", self.file_path)
            raise ParsingFailedError(
                f"Не удалось разобрать файл '{self.file_path.name}': {exc}"
            ) from exc
