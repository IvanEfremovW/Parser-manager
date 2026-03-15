"""
Фикстуры и конфигурация pytest для тестов Parser Manager.
"""

import io
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Callable

import pytest
from docx import Document as DocxDocument

from parser_manager.models import ParsedContent, DocumentMetadata, TextElement


# ─────────────────────────────────────────────────────────────────────────────
# Фикстуры временных директорий
# ─────────────────────────────────────────────────────────────────────────────


@pytest.fixture
def temp_dir() -> Path:
    """Создать временную директорию для тестовых файлов."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def temp_file(temp_dir: Path) -> Callable[[str, bytes], Path]:
    """Фабрика для создания временных файлов."""

    def _create_file(suffix: str, content: bytes) -> Path:
        file_path = temp_dir / f"test_file{suffix}"
        file_path.write_bytes(content)
        return file_path

    return _create_file


# ─────────────────────────────────────────────────────────────────────────────
# Фикстуры тестовых файлов для каждого формата
# ─────────────────────────────────────────────────────────────────────────────


@pytest.fixture
def sample_html_content() -> str:
    """Пример HTML контента для тестирования."""
    return """<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="author" content="Test Author">
    <meta name="description" content="Test Description">
    <title>Test Document Title</title>
</head>
<body>
    <h1>Main Heading</h1>
    <h2>Sub Heading</h2>
    <p>This is a test paragraph with some content.</p>
    <p>Another paragraph for testing.</p>
    <ul>
        <li>First list item</li>
        <li>Second list item</li>
    </ul>
    <table>
        <tr><td>Cell 1</td><td>Cell 2</td></tr>
        <tr><td>Cell 3</td><td>Cell 4</td></tr>
    </table>
    <a href="https://example.com">Test Link</a>
    <script>console.log("should be removed");</script>
    <style>.hidden { display: none; }</style>
</body>
</html>"""


@pytest.fixture
def sample_html_file(temp_dir: Path, sample_html_content: str) -> Path:
    """Создать пример HTML файла."""
    file_path = temp_dir / "test.html"
    file_path.write_text(sample_html_content, encoding="utf-8")
    return file_path


@pytest.fixture
def minimal_pdf_content() -> bytes:
    """Создать минимальный валидный PDF файл."""
    # Минимальный PDF с одной страницей и текстовым контентом
    text = "Test PDF Content"
    stream_content = f"BT /F1 12 Tf 50 700 Td ({text}) Tj ET".encode("latin-1", errors="replace")

    objects = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 5 0 R /Resources << /Font << /F1 4 0 R >> >> >>"
        ),
        4: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        5: (
            f"<< /Length {len(stream_content)} >>\nstream\n".encode()
            + stream_content
            + b"\nendstream"
        ),
    }

    buf = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for n in sorted(objects):
        offsets[n] = len(buf)
        buf += f"{n} 0 obj\n".encode() + objects[n] + b"\nendobj\n"

    xref_pos = len(buf)
    xref = f"xref\n0 {max(objects) + 1}\n0000000000 65535 f \n"
    for i in range(1, max(objects) + 1):
        xref += f"{offsets.get(i, 0):010d} 00000 {'n' if i in objects else 'f'} \n"

    trailer = f"trailer\n<< /Size {max(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n"
    buf += xref.encode() + trailer.encode()

    return bytes(buf)


@pytest.fixture
def sample_pdf_file(temp_dir: Path, minimal_pdf_content: bytes) -> Path:
    """Создать пример PDF файла."""
    file_path = temp_dir / "test.pdf"
    file_path.write_bytes(minimal_pdf_content)
    return file_path


@pytest.fixture
def sample_docx_file(temp_dir: Path) -> Path:
    """Создать пример DOCX файла с заголовками, параграфами и таблицами."""
    file_path = temp_dir / "test.docx"
    doc = DocxDocument()

    # Добавить метаданные
    doc.core_properties.title = "Test Document Title"
    doc.core_properties.author = "Test Author"
    doc.core_properties.subject = "Test Subject"
    doc.core_properties.comments = "Test Comments"

    # Добавить контент
    doc.add_heading("Main Heading", level=1)
    doc.add_heading("Sub Heading", level=2)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("Another paragraph for testing.")

    # Добавить список
    list_para = doc.add_paragraph("List item 1", style="List Bullet")
    list_para2 = doc.add_paragraph("List item 2", style="List Bullet")

    # Добавить таблицу
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Cell A1"
    table.rows[0].cells[1].text = "Cell B1"
    table.rows[1].cells[0].text = "Cell A2"
    table.rows[1].cells[1].text = "Cell B2"

    doc.save(str(file_path))
    return file_path


@pytest.fixture
def corrupted_file_content() -> bytes:
    """Контент, не являющийся валидным файлом документа."""
    return b"This is not a valid document file format"


@pytest.fixture
def sample_doc_file(temp_dir: Path, corrupted_file_content: bytes) -> Path:
    """Создать повреждённый DOC файл для тестов обработки ошибок."""
    file_path = temp_dir / "test.doc"
    file_path.write_bytes(corrupted_file_content)
    return file_path


@pytest.fixture
def sample_djvu_file(temp_dir: Path) -> Path:
    """Создать фейковый DJVU файл для тестов обработки ошибок."""
    file_path = temp_dir / "test.djvu"
    file_path.write_bytes(b"FAKE DJVU CONTENT")
    return file_path


# ─────────────────────────────────────────────────────────────────────────────
# Фикстуры моделей
# ─────────────────────────────────────────────────────────────────────────────


@pytest.fixture
def sample_metadata() -> DocumentMetadata:
    """Создать пример DocumentMetadata для тестирования."""
    return DocumentMetadata(
        title="Test Title",
        author="Test Author",
        subject="Test Subject",
        creation_date=datetime(2024, 1, 15, 10, 30, 0),
        modification_date=datetime(2024, 1, 16, 14, 45, 0),
        pages=10,
        language="en",
        encoding="UTF-8",
        custom_fields={"custom_field": "custom_value"},
    )


@pytest.fixture
def sample_text_elements() -> list[TextElement]:
    """Создать пример TextElements для тестирования."""
    return [
        TextElement(
            content="Main Heading",
            element_type="heading",
            level=1,
            page=1,
        ),
        TextElement(
            content="This is a paragraph.",
            element_type="paragraph",
            page=1,
        ),
        TextElement(
            content="Table content",
            element_type="table",
            page=2,
        ),
        TextElement(
            content="List item",
            element_type="list",
            page=2,
        ),
        TextElement(
            content="Link text",
            element_type="link",
            metadata={"href": "https://example.com"},
            page=3,
        ),
    ]


@pytest.fixture
def sample_parsed_content(sample_html_file: Path) -> ParsedContent:
    """Создать пример ParsedContent для тестирования."""
    return ParsedContent(
        file_path=str(sample_html_file),
        format="html",
        text="Main Heading\nSub Heading\nThis is a test paragraph.\nAnother paragraph.",
        metadata={
            "title": "Test Title",
            "author": "Test Author",
            "language": "en",
        },
        structure=[
            {"content": "Main Heading", "element_type": "heading", "level": 1},
            {"content": "Test paragraph", "element_type": "paragraph"},
        ],
        semantic_blocks=[
            {"content": "Main Heading", "element_type": "heading", "level": 1},
            {"content": "Test paragraph", "element_type": "paragraph"},
        ],
        quality={
            "overall_score": 0.85,
            "text_completeness": 0.9,
            "structure_score": 0.8,
            "noise_ratio": 0.05,
            "broken_chars_ratio": 0.01,
            "table_coverage": 0.1,
        },
        file_metrics={
            "file_name": "test.html",
            "file_size_bytes": 1024,
            "text_length": 100,
            "semantic_blocks": 2,
        },
        raw_data={"semantic_summary": {"total_blocks": 2}},
        success=True,
        error=None,
    )


# ─────────────────────────────────────────────────────────────────────────────
# Фикстуры API
# ─────────────────────────────────────────────────────────────────────────────


@pytest.fixture
def api_client():
    """Создать тестовый клиент для FastAPI приложения."""
    from fastapi.testclient import TestClient
    from parser_manager.api.app import app

    with TestClient(app) as client:
        yield client


@pytest.fixture
def sample_upload_file(sample_html_content: str):
    """Создать пример файла для загрузки."""
    from starlette.datastructures import UploadFile

    file_like = io.BytesIO(sample_html_content.encode("utf-8"))
    return UploadFile(filename="test.html", file=file_like)


# ─────────────────────────────────────────────────────────────────────────────
# Фикстура регистрации парсеров
# ─────────────────────────────────────────────────────────────────────────────


@pytest.fixture
def clean_parser_registry():
    """Фикстура для очистки реестра парсеров до и после тестов."""
    from parser_manager.core.parser_factory import ParserFactory

    # Сохранить текущий реестр
    saved_registry = ParserFactory._parsers_registry.copy()

    # Очистить реестр
    ParserFactory.clear_registry()

    yield ParserFactory

    # Восстановить реестр
    ParserFactory._parsers_registry = saved_registry
