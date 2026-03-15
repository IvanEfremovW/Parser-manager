"""
Unit tests for DocxParser.

Test Cases:
- TC-PARSER-DOCX-001: DOCX Parsing
"""

import pytest

from parser_manager.models import (
    CorruptedFileError,
    DocumentNotFoundError,
    ParsedContent,
)
from parser_manager.parsers.documents.docx_parser import DocxParser


class TestDocxParserInitialization:
    """Tests for DocxParser initialization."""

    def test_docx_parser_supported_extensions(self):
        """Test supported extensions."""
        assert DocxParser.supported_extensions == (".docx",)
        assert DocxParser.format_name == "docx"


class TestDocxParserParsing:
    """Tests for TC-PARSER-DOCX-001: DOCX Parsing."""

    def test_docx_parser_parse_success(self, sample_docx_file):
        """Test successful DOCX parsing."""
        parser = DocxParser(str(sample_docx_file))
        result = parser.parse()

        assert isinstance(result, ParsedContent)
        assert result.success is True
        assert result.format == "docx"
        assert result.error is None

    def test_docx_parser_extract_text(self, sample_docx_file):
        """Test text extraction from DOCX."""
        parser = DocxParser(str(sample_docx_file))
        text = parser.extract_text()

        assert isinstance(text, str)
        assert "Main Heading" in text
        assert "test paragraph" in text.lower()

    def test_docx_parser_extract_metadata(self, sample_docx_file):
        """Test metadata extraction from DOCX."""
        parser = DocxParser(str(sample_docx_file))
        metadata = parser.extract_metadata()

        assert metadata.title == "Test Document Title"
        assert metadata.author == "Test Author"
        assert metadata.subject == "Test Subject"

    def test_docx_parser_extract_structure(self, sample_docx_file):
        """Test structure extraction from DOCX."""
        parser = DocxParser(str(sample_docx_file))
        structure = parser.extract_structure()

        assert len(structure) > 0

        # Check for different element types
        element_types = [elem["element_type"] for elem in structure]

        assert "heading" in element_types
        assert "paragraph" in element_types

    def test_docx_parser_heading_levels(self, sample_docx_file):
        """Test that heading levels are correctly extracted."""
        parser = DocxParser(str(sample_docx_file))
        structure = parser.extract_structure()

        headings = [e for e in structure if e["element_type"] == "heading"]

        # Should have at least one heading
        assert len(headings) >= 1

        h1 = [h for h in headings if h["level"] == 1]
        h2 = [h for h in headings if h["level"] == 2]

        assert len(h1) >= 1
        assert len(h2) >= 1

    def test_docx_parser_table_extraction(self, sample_docx_file):
        """Test table extraction from DOCX."""
        parser = DocxParser(str(sample_docx_file))
        structure = parser.extract_structure()

        tables = [e for e in structure if e["element_type"] == "table"]

        assert len(tables) > 0
        # Table should contain cell content
        table_content = tables[0]["content"]
        assert "Cell A1" in table_content or "Cell" in table_content

    def test_docx_parser_list_extraction(self, sample_docx_file):
        """Test list extraction from DOCX."""
        parser = DocxParser(str(sample_docx_file))
        structure = parser.extract_structure()

        lists = [e for e in structure if e["element_type"] == "list"]

        # Should have list items
        assert len(lists) >= 1


class TestDocxParserEdgeCases:
    """Edge case tests for DocxParser."""

    def test_docx_parser_missing_file(self):
        """Test parsing non-existent file."""
        with pytest.raises(DocumentNotFoundError):
            DocxParser("/nonexistent/file.docx")

    def test_docx_parser_corrupted_file(self, temp_dir):
        """Test parsing corrupted DOCX file."""
        file_path = temp_dir / "corrupted.docx"
        file_path.write_bytes(b"Not a valid DOCX content")

        with pytest.raises(CorruptedFileError):
            parser = DocxParser(str(file_path))
            parser.extract_metadata()

    def test_docx_parser_empty_document(self, temp_dir):
        """Test parsing empty DOCX document."""
        from docx import Document

        file_path = temp_dir / "empty.docx"
        doc = Document()
        doc.save(str(file_path))

        parser = DocxParser(str(file_path))
        result = parser.parse()

        assert result.success is True
        assert result.text == ""

    def test_docx_parser_document_only_text(self, temp_dir):
        """Test parsing DOCX with only text content."""
        from docx import Document

        file_path = temp_dir / "text_only.docx"
        doc = Document()
        doc.add_paragraph("Just plain text.")
        doc.add_paragraph("Another paragraph.")
        doc.save(str(file_path))

        parser = DocxParser(str(file_path))
        result = parser.parse()

        assert result.success is True
        assert "Just plain text" in result.text
        assert "Another paragraph" in result.text

    def test_docx_parser_document_complex_structure(self, temp_dir):
        """Test parsing DOCX with complex structure."""
        from docx import Document

        file_path = temp_dir / "complex.docx"
        doc = Document()

        # Add various heading levels
        doc.add_heading("Title", level=0)
        doc.add_heading("Chapter 1", level=1)
        doc.add_heading("Section 1.1", level=2)
        doc.add_heading("Subsection 1.1.1", level=3)

        # Add paragraphs
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")

        # Add table
        table = doc.add_table(rows=3, cols=3)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"Cell {i},{j}"

        doc.save(str(file_path))

        parser = DocxParser(str(file_path))
        result = parser.parse()

        assert result.success is True

        structure = result.structure
        element_types = [e["element_type"] for e in structure]

        assert "heading" in element_types
        assert "paragraph" in element_types
        assert "table" in element_types

    def test_docx_parser_metadata_dates(self, temp_dir):
        """Test that metadata includes dates."""
        from datetime import datetime

        from docx import Document

        file_path = temp_dir / "dated.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.core_properties.created = datetime(2024, 1, 15, 10, 30, 0)
        doc.core_properties.modified = datetime(2024, 1, 16, 14, 45, 0)
        doc.save(str(file_path))

        parser = DocxParser(str(file_path))
        metadata = parser.extract_metadata()

        assert metadata.creation_date is not None or metadata.modification_date is not None

    def test_docx_parser_metadata_custom_fields(self, sample_docx_file):
        """Test that custom fields are included in metadata."""
        parser = DocxParser(str(sample_docx_file))
        metadata = parser.extract_metadata()

        metadata_dict = metadata.to_dict()

        # Check for custom fields
        assert "category" in metadata_dict or "comments" in metadata_dict

    def test_docx_parser_result_has_quality_metrics(self, sample_docx_file):
        """Test that parse result includes quality metrics."""
        parser = DocxParser(str(sample_docx_file))
        result = parser.parse()

        assert "quality" in result.to_dict()
        assert "overall_score" in result.quality
        assert 0.0 <= result.quality["overall_score"] <= 1.0

    def test_docx_parser_result_has_file_metrics(self, sample_docx_file):
        """Test that parse result includes file metrics."""
        parser = DocxParser(str(sample_docx_file))
        result = parser.parse()

        assert "file_metrics" in result.to_dict()
        assert result.file_metrics["file_name"] == sample_docx_file.name

    def test_docx_parser_result_has_semantic_blocks(self, sample_docx_file):
        """Test that parse result includes semantic blocks."""
        parser = DocxParser(str(sample_docx_file))
        result = parser.parse()

        assert len(result.semantic_blocks) > 0

    def test_docx_parser_result_has_raw_data(self, sample_docx_file):
        """Test that parse result includes raw_data."""
        parser = DocxParser(str(sample_docx_file))
        result = parser.parse()

        assert "raw_data" in result.to_dict()
        assert "semantic_summary" in result.raw_data

    def test_docx_parser_paragraphs_count(self, sample_docx_file):
        """Test that raw_data includes paragraph count."""
        parser = DocxParser(str(sample_docx_file))
        result = parser.parse()

        assert "paragraphs" in result.raw_data
        assert result.raw_data["paragraphs"] >= 0

    def test_docx_parser_style_name_none(self, temp_dir):
        """Test handling of paragraphs with None style name."""
        from docx import Document

        file_path = temp_dir / "no_style.docx"
        doc = Document()
        doc.add_paragraph("Test paragraph")
        # Some DOCX implementations may have None style
        doc.save(str(file_path))

        parser = DocxParser(str(file_path))
        result = parser.parse()

        assert result.success is True

    def test_docx_parser_unicode_content(self, temp_dir):
        """Test parsing DOCX with unicode content."""
        from docx import Document

        file_path = temp_dir / "unicode.docx"
        doc = Document()
        doc.add_heading("Заголовок", level=1)  # Russian
        doc.add_paragraph("Привет мир!")  # Russian
        doc.add_paragraph("你好世界!")  # Chinese
        doc.save(str(file_path))

        parser = DocxParser(str(file_path))
        result = parser.parse()

        assert result.success is True
        assert "Заголовок" in result.text
        assert "Привет мир!" in result.text
        assert "你好世界!" in result.text
